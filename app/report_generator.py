"""HTML report generator."""

import copy
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

from jinja2 import Environment, FileSystemLoader

from .config import settings
from .i18n import get_translator, _
from .models import AnalyzerResult, AuditResult, SeverityLevel


def translate_analyzer_content(result: AnalyzerResult, lang: str, translator) -> AnalyzerResult:
    """
    Translate analyzer result content to target language.

    This function handles translation at render time, allowing the analyzer
    code to remain in Ukrainian (source language) while supporting multiple
    output languages.
    """
    import re

    if lang == 'uk':
        return result  # Ukrainian is the source language

    # Create a deep copy to avoid modifying the original
    translated = copy.deepcopy(result)
    name = result.name

    # Translate theory
    theory_key = f"analyzer_content.{name}.theory"
    translated_theory = translator.get(theory_key, "")
    if translated_theory:
        translated.theory = translated_theory

    # Translate display_name
    display_name_key = f"analyzer_content.{name}.display_name"
    translated_display = translator.get(display_name_key, "")
    if translated_display:
        translated.display_name = translated_display

    # Translate description
    desc_key = f"analyzer_content.{name}.description"
    translated_desc = translator.get(desc_key, "")
    if translated_desc:
        translated.description = translated_desc

    # Translate summary - handle special cases for different analyzers
    if result.summary:
        if name == "cms":
            # Extract CMS name from Ukrainian summary
            cms_match = re.search(r'використовується (.+)$', result.summary)
            if cms_match:
                cms_name = cms_match.group(1)
                summary_key = f"analyzer_content.{name}.summary.cms_detected"
                translated_summary = translator.get(summary_key, "")
                if translated_summary and "{cms}" in translated_summary:
                    translated.summary = translated_summary.format(cms=cms_name)
            elif "не вдалося" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.cms_unknown"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

        elif name == "meta_tags":
            # Handle meta_tags summary with dynamic numbers
            if "Відсутні мета-теги:" in result.summary:
                # Extract numbers: "Відсутні мета-теги: X Title, Y Description"
                match = re.search(r'Відсутні мета-теги: (\d+) Title, (\d+) Description', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.missing"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(
                            missing_titles=match.group(1),
                            missing_descriptions=match.group(2)
                        )
                        # Handle duplicates part if present
                        dup_match = re.search(r'Дублікати: (\d+) Title, (\d+) Description', result.summary)
                        if dup_match:
                            dup_key = f"analyzer_content.{name}.summary.duplicates"
                            dup_trans = translator.get(dup_key, "")
                            if dup_trans:
                                translated.summary += ". " + dup_trans.format(
                                    duplicate_titles=dup_match.group(1),
                                    duplicate_descriptions=dup_match.group(2)
                                )
            elif "Дублікати:" in result.summary:
                dup_match = re.search(r'Дублікати: (\d+) Title, (\d+) Description', result.summary)
                if dup_match:
                    dup_key = f"analyzer_content.{name}.summary.duplicates"
                    translated_summary = translator.get(dup_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(
                            duplicate_titles=dup_match.group(1),
                            duplicate_descriptions=dup_match.group(2)
                        )
            elif "Всі" in result.summary and "мають коректні мета-теги" in result.summary:
                match = re.search(r'Всі (\d+) сторінок', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.all_ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(total_pages=match.group(1))

        elif name == "headings":
            if "Знайдено проблеми:" in result.summary:
                match = re.search(r'Знайдено проблеми: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems_found"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(problems=match.group(1))
            elif "мають коректний H1" in result.summary:
                match = re.search(r'Всі (\d+) сторінок', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.all_ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(total_pages=match.group(1))

        elif name == "page_404":
            summary_map = {
                "налаштована коректно": "ok",
                "створити або виправити": "missing",
                "потребує покращень": "needs_improvement",
                "Не вдалося перевірити": "check_failed"
            }
            for ukr_text, key in summary_map.items():
                if ukr_text in result.summary:
                    summary_key = f"analyzer_content.{name}.summary.{key}"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary
                    break

        elif name == "speed":
            # Extract mobile and desktop scores
            match = re.search(r'Mobile: (\d+)/100, Desktop: (\d+)/100', result.summary)
            if match:
                mobile, desktop = match.group(1), match.group(2)
                if "в нормі" in result.summary:
                    key = "ok"
                elif "Потрібна оптимізація" in result.summary:
                    key = "needs_optimization"
                elif "Критичні" in result.summary:
                    key = "critical"
                else:
                    key = None
                if key:
                    summary_key = f"analyzer_content.{name}.summary.{key}"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(mobile=mobile, desktop=desktop)
            elif "Не вдалося" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.failed"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

        elif name == "images":
            if "Всі" in result.summary and "оптимізовані" in result.summary:
                match = re.search(r'Всі (\d+) зображень', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.all_ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(count=match.group(1))
            elif "Знайдено" in result.summary:
                match = re.search(r'Знайдено (\d+) зображень\. Проблеми: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(count=match.group(1), problems=match.group(2))

        elif name == "content":
            if "мають достатньо контенту" in result.summary:
                match = re.search(r'Всі (\d+) сторінок.*Середня кількість слів: (\d+)', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.all_ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(total_pages=match.group(1), avg_words=match.group(2))
            elif "Проблеми з контентом:" in result.summary:
                match = re.search(r'Проблеми з контентом: (.+)\. Середня кількість слів: (\d+)', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(problems=match.group(1), avg_words=match.group(2))

        elif name == "links":
            if "Проблем не знайдено" in result.summary:
                match = re.search(r'Перевірено (\d+) внутрішніх та (\d+) зовнішніх', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.no_broken"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(internal=match.group(1), external=match.group(2))
            elif "Знайдено битих посилань:" in result.summary:
                match = re.search(r'Знайдено битих посилань: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.broken_found"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(broken=match.group(1))

        elif name == "favicon":
            summary_map = {
                "налаштовано коректно": "ok",
                "відсутній": "missing",
                "можна покращити": "needs_improvement"
            }
            for ukr_text, key in summary_map.items():
                if ukr_text in result.summary:
                    summary_key = f"analyzer_content.{name}.summary.{key}"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary
                    break

        elif name == "external_links":
            match = re.search(r'Знайдено (\d+) зовнішніх посилань на (\d+) доменів', result.summary)
            if match:
                summary_key = f"analyzer_content.{name}.summary.ok"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary.format(count=match.group(1), domains=match.group(2))
            else:
                match = re.search(r'Знайдено (\d+) зовнішніх посилань\. Попереджень: (\d+), інфо: (\d+)', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.with_warnings"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(count=match.group(1), warnings=match.group(2), info=match.group(3))

        elif name == "robots":
            if "в порядку" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.ok"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary
            elif "Проблеми з індексацією:" in result.summary:
                match = re.search(r'Проблеми з індексацією: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(problems=match.group(1))

        elif name == "structure":
            if "оптимальна" in result.summary:
                match = re.search(r'Максимальна глибина: (\d+) рівнів', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(depth=match.group(1))
            else:
                match = re.search(r'Максимальна глибина: (\d+)\. Проблеми: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(depth=match.group(1), problems=match.group(2))

        elif name == "content_sections":
            if "Виявлено:" in result.summary:
                match = re.search(r'Виявлено: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.detected"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(sections=match.group(1))
            elif "не виявлено" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.not_detected"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

    # Translate issues
    for issue in translated.issues:
        # Try to translate message by category
        msg_key = f"analyzer_content.{name}.issues.{issue.category}"
        translated_msg = translator.get(msg_key, "")

        if translated_msg:
            try:
                # Handle CMS-specific translations with dynamic CMS name
                if name == "cms" and issue.category == "cms_detected":
                    # Extract CMS name from original Ukrainian message
                    cms_match = re.search(r'використовується (.+)$', issue.message)
                    if cms_match and "{cms}" in translated_msg:
                        cms_name = cms_match.group(1)
                        issue.message = translated_msg.format(cms=cms_name)
                elif name == "cms" and issue.category == "multiple_cms":
                    # Extract CMS list from original message
                    cms_match = re.search(r'ознаки: (.+)$', issue.message)
                    if cms_match and "{cms_list}" in translated_msg:
                        cms_list = cms_match.group(1)
                        issue.message = translated_msg.format(cms_list=cms_list)
                # Try to format with count if available
                elif issue.count is not None and "{count}" in translated_msg:
                    issue.message = translated_msg.format(count=issue.count)
                elif "{" not in translated_msg:
                    # No placeholders, use as-is
                    issue.message = translated_msg
                # If has other placeholders, keep original message
            except (KeyError, ValueError):
                # If formatting fails, keep original message
                pass

        # Translate details - handle CMS special case
        details_key = f"analyzer_content.{name}.details.{issue.category}"
        translated_details = translator.get(details_key, "")
        if translated_details:
            if name == "cms" and issue.category == "cms_detected":
                # Extract evidence from original details
                evidence_match = re.search(r'ознаки: (.+)$', issue.details or "")
                if evidence_match and "{evidence}" in translated_details:
                    evidence = evidence_match.group(1)
                    issue.details = translated_details.format(evidence=evidence)
            elif "{" not in translated_details:
                issue.details = translated_details

        # Translate recommendation
        rec_key = f"analyzer_content.{name}.recommendations.{issue.category}"
        translated_rec = translator.get(rec_key, "")
        if translated_rec and "{" not in translated_rec:
            issue.recommendation = translated_rec

    # Translate tables
    table_titles = translator.translations.get("table_translations", {}).get("titles", {})
    table_headers = translator.translations.get("table_translations", {}).get("headers", {})

    for table in translated.tables:
        # Translate table title
        if table.get("title") and table["title"] in table_titles:
            table["title"] = table_titles[table["title"]]

        # Translate table headers
        if table.get("headers"):
            table["headers"] = [
                table_headers.get(h, h) for h in table["headers"]
            ]

        # Translate row values that are in the headers translation map
        if table.get("rows"):
            new_rows = []
            for row in table["rows"]:
                new_row = {}
                for key, value in row.items():
                    # Translate key if it's in table_headers
                    new_key = table_headers.get(key, key)
                    # Translate value if it's a simple string in table_headers
                    if isinstance(value, str) and value in table_headers:
                        new_row[new_key] = table_headers[value]
                    else:
                        new_row[new_key] = value
                new_rows.append(new_row)
            table["rows"] = new_rows

    return translated


class ReportGenerator:
    """Generates autonomous HTML reports from audit results."""

    def __init__(self):
        templates_path = Path(__file__).parent / "templates"
        self.env = Environment(
            loader=FileSystemLoader(str(templates_path)),
            autoescape=True,
        )

        # Add custom filters
        self.env.filters['status_icon'] = self.status_icon
        self.env.filters['severity_class'] = self.severity_class
        self.env.filters['format_number'] = self.format_number

    @staticmethod
    def status_icon(severity: SeverityLevel) -> str:
        """Convert severity to status icon."""
        icons = {
            SeverityLevel.SUCCESS: '✓',
            SeverityLevel.WARNING: '⚠️',
            SeverityLevel.ERROR: '✗',
            SeverityLevel.INFO: 'ℹ️',
        }
        return icons.get(severity, '')

    @staticmethod
    def severity_class(severity: SeverityLevel) -> str:
        """Convert severity to CSS class."""
        classes = {
            SeverityLevel.SUCCESS: 'success',
            SeverityLevel.WARNING: 'warning',
            SeverityLevel.ERROR: 'error',
            SeverityLevel.INFO: 'info',
        }
        return classes.get(severity, 'info')

    @staticmethod
    def format_number(value: int) -> str:
        """Format number with thousands separator."""
        return f"{value:,}".replace(",", " ")

    async def generate(self, audit: AuditResult) -> str:
        """Generate HTML report and return file path."""
        template = self.env.get_template("report.html")

        # Get translator for the audit language
        lang = getattr(audit, 'language', 'uk') or 'uk'
        t = get_translator(lang)

        # Prepare sections for navigation with translated names
        sections = []
        section_order = [
            ("cms", "🔧"),
            ("meta_tags", "🏷️"),
            ("headings", "📝"),
            ("page_404", "🚫"),
            ("speed", "⚡"),
            ("images", "🖼️"),
            ("content", "📄"),
            ("links", "🔗"),
            ("favicon", "🌟"),
            ("external_links", "🔗"),
            ("robots", "🤖"),
            ("structure", "🏗️"),
            ("content_sections", "📰"),
        ]

        for name, icon in section_order:
            if name in audit.results:
                result = audit.results[name]

                # Translate analyzer content if not Ukrainian
                if lang != 'uk':
                    result = translate_analyzer_content(result, lang, t)

                # Get translated title, fallback to display_name from result
                title = t(f"analyzers.{name}.name")
                if title == f"analyzers.{name}.name":
                    title = result.display_name  # Fallback to analyzer's display_name

                sections.append({
                    "id": name,
                    "title": title,
                    "icon": icon,
                    "severity": result.severity,
                    "result": result,
                })

        # Extract domain
        domain = urlparse(audit.url).netloc.replace("www.", "")

        # Prepare translations for template
        translations = {
            "report_title": t("report.title"),
            "overview": t("report.overview"),
            "pages_crawled": t("report.pages_crawled"),
            "passed_checks": t("report.passed_checks"),
            "warnings": t("report.warnings"),
            "critical_issues": t("report.critical_issues"),
            "theory_title": t("report.theory_title"),
            "examples": t("report.examples"),
            "recommendation": t("report.recommendation"),
            "no_issues": t("report.no_issues"),
            "expand_more": t("common.expand_more"),
            "collapse": t("common.collapse"),
            "pagespeed_screenshots": t("report.pagespeed_screenshots"),
        }

        # Render template
        html = template.render(
            audit=audit,
            domain=domain,
            sections=sections,
            generated_at=datetime.now().strftime("%d.%m.%Y %H:%M"),
            SeverityLevel=SeverityLevel,
            t=translations,
            lang=lang,
        )

        # Save report
        report_filename = f"audit_{audit.id}.html"
        report_path = Path(settings.REPORTS_DIR) / report_filename

        with open(report_path, "w", encoding="utf-8") as f:
            f.write(html)

        return str(report_path)

    async def generate_pdf(self, audit: AuditResult) -> str:
        """Generate PDF report and return file path."""
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            raise ImportError("weasyprint is required for PDF export. Install it with: pip install weasyprint")

        # First generate HTML
        html_path = await self.generate(audit)

        # Read the HTML content
        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        # Create PDF
        pdf_filename = f"audit_{audit.id}.pdf"
        pdf_path = Path(settings.REPORTS_DIR) / pdf_filename

        # Add print-specific CSS to hide sidebar and adjust layout
        print_css = CSS(string="""
            @page {
                size: A4;
                margin: 1.5cm;
            }
            .sidebar {
                display: none !important;
            }
            .main {
                margin-left: 0 !important;
                padding: 0 !important;
                max-width: 100% !important;
            }
            body {
                background: white !important;
                font-size: 11pt !important;
            }
            .section {
                page-break-inside: avoid;
                break-inside: avoid;
            }
            .summary-grid {
                grid-template-columns: repeat(4, 1fr) !important;
            }
            .screenshots-grid {
                grid-template-columns: 1fr !important;
            }
            .screenshot-card img {
                max-width: 100% !important;
            }
            .theory-block {
                page-break-inside: avoid;
            }
            /* Force details open and hide interactive elements for PDF */
            details {
                display: block !important;
            }
            details summary {
                list-style: none !important;
                pointer-events: none !important;
            }
            details summary::marker,
            details summary::-webkit-details-marker {
                display: none !important;
            }
            details:not([open]) > *:not(summary) {
                display: block !important;
            }
            .expand-urls-btn {
                display: none !important;
            }
            .urls-hidden {
                display: block !important;
            }
        """)

        # Force all details elements to be open for PDF
        html_content = html_content.replace('<details class="theory-block">', '<details class="theory-block" open>')

        HTML(string=html_content).write_pdf(pdf_path, stylesheets=[print_css])

        return str(pdf_path)

    # --- DOCX Helper Methods ---

    @staticmethod
    def _docx_set_cell_shading(cell, color_hex: str):
        """Set background color on a table cell."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        tc_pr.append(shading)

    @staticmethod
    def _docx_set_cell_left_border(cell, color_hex: str, width: str = "24"):
        """Add a colored left border to a cell."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tc_pr.append(borders)
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), width)
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), color_hex)
        borders.append(left)

    @staticmethod
    def _docx_remove_cell_borders(cell):
        """Remove all borders from a cell."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for side in ('w:top', 'w:left', 'w:bottom', 'w:right'):
            el = OxmlElement(side)
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
            borders.append(el)
        tc_pr.append(borders)

    @staticmethod
    def _docx_set_font(run, font_name: str = 'Inter', size_pt=None, bold=None, color_rgb=None):
        """Configure a run with font settings."""
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        run.font.name = font_name
        r = run._element
        r_pr = r.find(qn('w:rPr'))
        if r_pr is None:
            r_pr = r.makeelement(qn('w:rPr'), {})
            r.insert(0, r_pr)
        r_fonts = r_pr.find(qn('w:rFonts'))
        if r_fonts is None:
            r_fonts = r.makeelement(qn('w:rFonts'), {})
            r_pr.append(r_fonts)
        r_fonts.set(qn('w:ascii'), font_name)
        r_fonts.set(qn('w:hAnsi'), font_name)
        r_fonts.set(qn('w:cs'), font_name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        if color_rgb is not None:
            run.font.color.rgb = RGBColor(*color_rgb)

    def _docx_parse_theory(self, doc, theory_html: str):
        """Parse theory HTML into Word paragraphs with formatting."""
        import re
        from docx.shared import Pt, RGBColor

        if not theory_html:
            return

        # Create a single-cell table for gray background
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        self._docx_remove_cell_borders(cell)
        self._docx_set_cell_shading(cell, 'F0F4F8')

        # Clear default paragraph
        cell.text = ''

        lines = theory_html.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            p = cell.add_paragraph() if cell.paragraphs[0].text or i > 0 else cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Handle bullet points
            if line.startswith('•'):
                line = line[1:].strip()
                # Add bullet character
                run = p.add_run('  •  ')
                self._docx_set_font(run, size_pt=9, color_rgb=(107, 114, 128))

            # Parse inline HTML tags
            parts = re.split(r'(<strong>.*?</strong>|<code>.*?</code>)', line)
            for part in parts:
                if not part:
                    continue
                strong_match = re.match(r'<strong>(.*?)</strong>', part)
                code_match = re.match(r'<code>(.*?)</code>', part)
                if strong_match:
                    run = p.add_run(strong_match.group(1))
                    self._docx_set_font(run, size_pt=9, bold=True)
                elif code_match:
                    run = p.add_run(code_match.group(1))
                    self._docx_set_font(run, font_name='Consolas', size_pt=9, color_rgb=(107, 114, 128))
                else:
                    # Strip any remaining HTML
                    clean = re.sub(r'<[^>]+>', '', part)
                    if clean:
                        run = p.add_run(clean)
                        self._docx_set_font(run, size_pt=9)

    def _docx_add_issue_card(self, doc, issue, t_labels: dict):
        """Add a colored issue card as a borderless single-cell table."""
        from docx.shared import Pt, RGBColor

        severity_colors = {
            SeverityLevel.ERROR: 'FEE2E2',
            SeverityLevel.WARNING: 'FEF3C7',
            SeverityLevel.SUCCESS: 'D1FAE5',
            SeverityLevel.INFO: 'DBEAFE',
        }
        severity_text_colors = {
            SeverityLevel.ERROR: (239, 68, 68),
            SeverityLevel.WARNING: (180, 120, 0),
            SeverityLevel.SUCCESS: (16, 150, 100),
            SeverityLevel.INFO: (59, 130, 246),
        }
        severity_icons = {
            SeverityLevel.SUCCESS: "✓",
            SeverityLevel.WARNING: "⚠",
            SeverityLevel.ERROR: "✗",
            SeverityLevel.INFO: "ℹ",
        }

        bg_color = severity_colors.get(issue.severity, 'F3F4F6')
        text_color = severity_text_colors.get(issue.severity, (31, 41, 55))
        icon = severity_icons.get(issue.severity, "")

        # Create single-cell table for the card
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        self._docx_remove_cell_borders(cell)
        self._docx_set_cell_shading(cell, bg_color)

        # Issue message (bold, colored)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{icon} {issue.message}")
        self._docx_set_font(run, size_pt=10, bold=True, color_rgb=text_color)

        # Details
        if issue.details:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(issue.details)
            self._docx_set_font(run, size_pt=9, color_rgb=(55, 65, 81))

        # Recommendation
        if issue.recommendation:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            rec_label = t_labels.get("recommendation", "Рекомендація")
            run = p.add_run(f"💡 {rec_label}: ")
            self._docx_set_font(run, size_pt=9, bold=True, color_rgb=(55, 65, 81))
            run = p.add_run(issue.recommendation)
            self._docx_set_font(run, size_pt=9, color_rgb=(55, 65, 81))

        # Affected URLs
        if issue.affected_urls:
            examples_label = t_labels.get("examples", "Приклади")
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(f"{examples_label}:")
            self._docx_set_font(run, size_pt=8, bold=True, color_rgb=(107, 114, 128))
            for url in issue.affected_urls[:5]:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(f"  • {url}")
                self._docx_set_font(run, size_pt=8, color_rgb=(107, 114, 128))

    async def generate_docx(self, audit: AuditResult) -> str:
        """Generate styled DOCX report and return file path."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError("python-docx is required for Word export. Install it with: pip install python-docx")

        # Setup i18n
        lang = getattr(audit, 'language', 'uk') or 'uk'
        t = get_translator(lang)

        t_labels = {
            "express_title": t("report.express_title"),
            "generated_at": t("report.generated_at"),
            "overview": t("report.overview"),
            "pages_crawled": t("report.pages_crawled"),
            "passed_checks": t("report.passed_checks"),
            "warnings": t("report.warnings"),
            "critical_issues": t("report.critical_issues"),
            "theory_title": t("report.theory_title"),
            "examples": t("report.examples"),
            "recommendation": t("report.recommendation"),
            "no_issues": t("report.no_issues"),
        }

        # Extract domain
        domain = urlparse(audit.url).netloc.replace("www.", "")

        # Create document
        doc = Document()

        # --- Setup Inter font for styles ---
        from docx.oxml import OxmlElement

        def _set_style_font(s, font_name='Inter'):
            s.font.name = font_name
            rPr = s.element.get_or_add_rPr()
            r_fonts = rPr.find(qn('w:rFonts'))
            if r_fonts is None:
                r_fonts = OxmlElement('w:rFonts')
                rPr.append(r_fonts)
            r_fonts.set(qn('w:ascii'), font_name)
            r_fonts.set(qn('w:hAnsi'), font_name)
            r_fonts.set(qn('w:cs'), font_name)

        style = doc.styles['Normal']
        style.font.size = Pt(10)
        _set_style_font(style)

        for heading_level in range(1, 4):
            style_name = f'Heading {heading_level}'
            if style_name in doc.styles:
                h_style = doc.styles[style_name]
                h_style.font.color.rgb = RGBColor(31, 41, 55)
                _set_style_font(h_style)

        # --- Title ---
        title_text = f"{t_labels['express_title']}: {domain}"
        title_para = doc.add_heading(title_text, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            self._docx_set_font(run, size_pt=22, bold=True, color_rgb=(31, 41, 55))

        # Subtitle with date
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{t_labels['generated_at']}: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        self._docx_set_font(run, size_pt=11, color_rgb=(128, 128, 128))

        # URL
        url_para = doc.add_paragraph()
        url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = url_para.add_run(audit.url)
        self._docx_set_font(run, size_pt=10, color_rgb=(59, 130, 246))

        doc.add_paragraph()

        # --- Summary Section ---
        doc.add_heading(f"📊 {t_labels['overview']}", level=1)

        summary_table = doc.add_table(rows=2, cols=4)
        summary_table.style = 'Table Grid'

        summary_items = [
            (t_labels['pages_crawled'], str(audit.pages_crawled), '3B82F6'),
            (t_labels['passed_checks'], str(audit.passed_checks), '10B981'),
            (t_labels['warnings'], str(audit.warnings), 'F59E0B'),
            (t_labels['critical_issues'], str(audit.critical_issues), 'EF4444'),
        ]

        for i, (label, value, color) in enumerate(summary_items):
            # Header cell
            header_cell = summary_table.rows[0].cells[i]
            header_cell.text = ''
            p = header_cell.paragraphs[0]
            run = p.add_run(label)
            self._docx_set_font(run, size_pt=9, bold=True, color_rgb=(107, 114, 128))
            self._docx_set_cell_left_border(header_cell, color, '24')

            # Value cell
            value_cell = summary_table.rows[1].cells[i]
            value_cell.text = ''
            p = value_cell.paragraphs[0]
            run = p.add_run(value)
            self._docx_set_font(run, size_pt=18, bold=True, color_rgb=(31, 41, 55))
            self._docx_set_cell_left_border(value_cell, color, '24')

        doc.add_paragraph()

        # --- Results Sections ---
        section_order = [
            ("cms", "🔧"),
            ("meta_tags", "🏷️"),
            ("headings", "📝"),
            ("page_404", "🚫"),
            ("speed", "⚡"),
            ("images", "🖼️"),
            ("content", "📄"),
            ("links", "🔗"),
            ("favicon", "🌟"),
            ("external_links", "🔗"),
            ("robots", "🤖"),
            ("structure", "🏗️"),
            ("content_sections", "📰"),
        ]

        severity_badge_text = {
            SeverityLevel.SUCCESS: ("✓", (16, 185, 129)),
            SeverityLevel.WARNING: ("⚠", (245, 158, 11)),
            SeverityLevel.ERROR: ("✗", (239, 68, 68)),
            SeverityLevel.INFO: ("ℹ", (59, 130, 246)),
        }

        for name, icon in section_order:
            if name not in audit.results:
                continue

            result = audit.results[name]

            # Translate content if needed
            if lang != 'uk':
                result = translate_analyzer_content(result, lang, t)

            # Get translated section title
            section_title = t(f"analyzers.{name}.name")
            if section_title == f"analyzers.{name}.name":
                section_title = result.display_name

            # Section heading with emoji
            heading = doc.add_heading(f"{icon} {section_title}", level=1)

            # Add severity badge after heading
            badge_text, badge_color = severity_badge_text.get(
                result.severity, ("ℹ", (59, 130, 246))
            )
            run = heading.add_run(f"  [{badge_text}]")
            self._docx_set_font(run, size_pt=12, bold=False, color_rgb=badge_color)

            # Summary
            if result.summary:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(result.summary)
                self._docx_set_font(run, size_pt=10, bold=True)

            # Theory section
            if result.theory:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(f"📖 {t_labels['theory_title']}")
                self._docx_set_font(run, size_pt=10, bold=True, color_rgb=(75, 85, 99))
                self._docx_parse_theory(doc, result.theory)
                doc.add_paragraph()  # spacing after theory

            # Issues
            if result.issues:
                for issue in result.issues:
                    self._docx_add_issue_card(doc, issue, t_labels)
                    # Small spacing between cards
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(2)
                    spacer.paragraph_format.space_after = Pt(2)
            elif not result.tables:
                p = doc.add_paragraph()
                run = p.add_run(f"✅ {t_labels['no_issues']}")
                self._docx_set_font(run, size_pt=10, color_rgb=(16, 185, 129))

            # Tables
            for table_info in result.tables:
                if not table_info.get("rows"):
                    continue

                # Table title
                table_title = table_info.get("title", "")
                if table_title:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(8)
                    run = p.add_run(table_title)
                    self._docx_set_font(run, size_pt=10, bold=True)

                headers = table_info.get("headers", [])
                rows = table_info.get("rows", [])

                if headers and rows:
                    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    table.style = 'Table Grid'

                    # Header row with gray background
                    for col_idx, header in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(header)
                        self._docx_set_font(run, size_pt=9, bold=True)
                        self._docx_set_cell_shading(cell, 'F3F4F6')

                    # Data rows
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, header in enumerate(headers):
                            value = row_data.get(header, "")
                            cell = table.rows[row_idx + 1].cells[col_idx]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            run = p.add_run(str(value))
                            self._docx_set_font(run, size_pt=9)
                            # Alternating row shading
                            if row_idx % 2 == 1:
                                self._docx_set_cell_shading(cell, 'F9FAFB')

            doc.add_paragraph()  # spacing between sections

        # Save document
        docx_filename = f"audit_{audit.id}.docx"
        docx_path = Path(settings.REPORTS_DIR) / docx_filename
        doc.save(docx_path)

        return str(docx_path)
